from docx import Document
import re
def replace_str(original_string):
    tstr= original_string.replace(' ', '').replace('\t', '').\
        replace('【', '[').replace('】', ']').replace('\n','').replace('\r','').\
        replace('：', ':').replace("（", "(").replace("）", ")").replace('\r\n','').replace('、','')


    new_str1=re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]',"", tstr)
    new_str = re.sub(r'[\n\r]', ' ',new_str1 )
    return new_str
def extract_tables_with_keys_from_docx(file_path):
    doc = Document(file_path)

    tables_with_keys = {}
    table_index = 0
    for paragraph in doc.paragraphs:
        # 检查段落之后是否紧跟着表格
        if table_index < len(doc.tables) and paragraph._element.getnext() is doc.tables[table_index]._element:
            # 使用段落文本作为键，读取紧跟其后的表格作为值
            key = paragraph.text
            table = doc.tables[table_index]
            table_content = []
            for row in table.rows:
                for cell in row.cells:
                    # 去除单元格文本中的空格和换行符
                    cleaned_text = replace_str(cell.text)
                    table_content.append(cleaned_text)
                table_content.append(table_content)
            tables_with_keys[key] = table_content
            table_index += 1

    return tables_with_keys

# 调用函数并传入你的Word文档的路径
tables_with_keys = extract_tables_with_keys_from_docx( r"z:\公司差旅费管理实施细则（咨财〔2019〕2583号，2019年修订）.doc")
for key, table in tables_with_keys.items():
    print(f"Table - {key}:{table}")
    # if "111磊大大大" in table:
    #     print("ok")
    # for row in table:
    #     print(row)
