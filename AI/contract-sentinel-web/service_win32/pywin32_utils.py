import tempfile
import platform
from imessagefilter import CMessageFilter
import win32com.client as win32
import time, os, copy, shutil
from service.utils import replace_str
import subprocess, sys
from service.g import logger
from tools.utils import *
from docx import Document
from service.utils import replace_str

# 获取 Python 解释器的位数
bits = platform.architecture()[0]

# 判断是否为 32 位
is_32bit = bits == '32bit'

def close_word_window():
    # os.system("chcp 65001")
    # 执行命令，不显示输出和错误
    subprocess.run("chcp 65001", stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, shell=True)
    subprocess.run("taskkill /F /IM WINWORD.EXE", shell=True, text=True)


class WordMix():

    def message_filter_register(self):
        if is_32bit:
            CMessageFilter.register()
        else:
            pass

    def message_filter_revoke(self):
        if is_32bit:
            CMessageFilter.revoke()
        else:
            pass

    def word_app_dispatch(self):
        self.word = win32.gencache.EnsureDispatch("Word.Application")
        # print(sys.modules[self.word.__module__].__file__)
        self.word.Visible = False
        return self.word

    def word_app_quit(self, word):
        word.Application.Quit()

    def request_word(self, force_close=False):
        if force_close:
            close_word_window()
        self.message_filter_register()
        word = self.word_app_dispatch()
        return word

    def release_word(self, word, doc, doc_save=True, force_close=False):
        if doc is not None:
            doc.Close(doc_save)
        # 关闭Word应用程序
        self.word_app_quit(word)
        self.message_filter_revoke()
        if force_close:
            close_word_window()

    def doc_parse(self, path_file_name):
        word = self.request_word(force_close=True)
        try:
            doc = word.Documents.Open(path_file_name)
            doc.Activate()

        except Exception as e:
            return None, None, e

        return word, doc, None

    # 保留特殊字符 to txt
    def doc_to_txt(self, path, txt_path, force_close=False):
        word, doc, e = self.doc_parse(path)
        try:
            if e is not None:
                raise e
            # Extract text from the document
            text_content = doc.Content.Text
            # Write the extracted text to a text file
            with open(txt_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(text_content)
        finally:
            self.release_word(word, doc, force_close=True)

    def custom_doc_to_txt(self, path, txt_path, force_close=False):
        word, doc, e = self.doc_parse(path)
        try:
            if e is not None:
                raise e
            text_dict = {}
            text_dict_map = {}
            # Extract text from the document
            paragraphs = doc.Paragraphs

            with open(txt_path, 'w', encoding='utf-8') as txt_file:
                for paragraph in paragraphs:
                    text = copy.deepcopy(paragraph.Range.Text)
                    txt_file.write(replace_str(text))
                    txt_file.write('\n')

                    section_number = paragraph.Range.ListFormat.ListString
                    start_page_number = paragraph.Range.Information(win32.constants.wdActiveEndPageNumber)
                    # 获取段落的样式
                    # paragraph_style = paragraph.Range.ParagraphFormat.Style.NameLocal

                    # print(f"当前文本:{replace_str(text)}")
                    # print(f"当前段落序号:{section_number}")
                    # print(f"当前页码:{start_page_number}")
                    text_dict[replace_str(text)] = {"page": replace_str(f"{start_page_number}"),
                                                    "part": replace_str(section_number)}
                    text_dict_map[replace_str(text)] = text

        finally:
            self.release_word(word, doc, force_close=True)
        return text_dict, text_dict_map

    def create_revisions_docx(self, doc_path1, doc_path2, output_path):
        word = self.request_word()
        try:
            # 不对比 格式 ，空格
            word.CompareDocuments(word.Documents.Open(doc_path1),
                                  word.Documents.Open(doc_path2), CompareFormatting=False, CompareWhitespace=False,
                                  IgnoreAllComparisonWarnings=True)

            # 引用常量模块
            word_constants = win32.constants
            # word.ActiveDocument.ActiveWindow.View.Type = 3
            # 设置文档为显示修订模式
            word.ActiveDocument.TrackRevisions = True

            # 显示所有修订标记
            word.ActiveWindow.View.RevisionsFilter.Markup = word_constants.wdRevisionsMarkupAll  # 1 表示显示所有标记

            # 显示所有标记，包括格式更改
            word.ActiveWindow.View.RevisionsFilter.View = word_constants.wdRevisionsViewFinal
            #
            word.ActiveDocument.SaveAs(FileName=output_path)
        finally:
            self.release_word(word, None, force_close=True)

    def table_output_list(self,doc_path):
        doc = Document(doc_path)

        tables_with_keys = {}
        table_index = 0
        for paragraph in doc.paragraphs:
            # 检查段落之后是否紧跟着表格
            if table_index < len(doc.tables) and paragraph._element.getnext() is doc.tables[table_index]._element:
                # 使用段落文本作为键，读取紧跟其后的表格作为值
                key = paragraph.text
                table = doc.tables[table_index]
                table_content = []
                for row in table.rows:
                    for cell in row.cells:
                        # 去除单元格文本中的空格和换行符
                        cleaned_text = replace_str(cell.text)
                        table_content.append(cleaned_text)
                    table_content.append(table_content)
                tables_with_keys[key] = table_content
                table_index += 1

        return tables_with_keys

    def revision_stat(self, path_file_name, save_file, accept=False):
        word, doc, e = self.doc_parse(path_file_name)
        try:
            self.word.ActiveDocument.TrackRevisions = True  # 关闭默认修订模式 Maybe not need this (not really but why not)

            # Accept all revisions
            if accept:
                self.word.ActiveDocument.Revisions.AcceptAll()
            else:
                self.word.ActiveDocument.Revisions.RejectAll()
            # Delete all comments
            if self.word.ActiveDocument.Comments.Count >= 1:
                self.word.ActiveDocument.DeleteAllComments()

            word.ActiveDocument.SaveAs(FileName=save_file)
        finally:

            self.release_word(word, doc, force_close=True)

    def revisions_index_extract_and_build_page(self, doc_path, comments_dict, text_dict, accept_text_dict):
        word, doc, e = self.doc_parse(doc_path)
        revision_index = {}
        src_text_dict = {}
        dest_text_dict = {}
        try:
            for index, paragraph in enumerate(doc.Paragraphs):
                revisions = paragraph.Range.Revisions
                section_number = paragraph.Range.ListFormat.ListString
                start_page_number = paragraph.Range.Information(win32.constants.wdActiveEndPageNumber)
                page = replace_str(f"{start_page_number}")
                part = replace_str(section_number)

                src_text = replace_str(copy.deepcopy(paragraph.Range.Text))
                if revisions.Count > 0:
                    revisions.AcceptAll()
                dest_text = replace_str(copy.deepcopy(paragraph.Range.Text))

                # 表格中的文本改动后paragraph.Range.Revisions.Count 无法判断，需要table.Range.Revisions.Count，见pywin32_表格

                # xxx 默认使用邮件word中的页码
                if src_text in comments_dict.keys():
                    src_text_dict.setdefault(src_text, {"page": page, "part": part, "index": index})

                elif dest_text in comments_dict.keys():
                    dest_text_dict.setdefault(dest_text, {"page": page, "part": part, "index": index})

            for key, value in src_text_dict.items():
                if key in comments_dict.keys():
                    revision_index.setdefault(value['index'], comments_dict[key])
                    text_dict[key]["page"] = value['page']
                    text_dict[key]["part"] = value['part']
                    # if key in accept_text_dict:
                    #     text_dict[key]["page"]=accept_text_dict[key]['page']
                    #     text_dict[key]["part"]=accept_text_dict[key]['part']

            for key, value in dest_text_dict.items():
                if key in comments_dict.keys() and key not in src_text_dict.keys():
                    revision_index.setdefault(value['index'], comments_dict[key])
                    text_dict[key]["page"] = value['page']
                    text_dict[key]["part"] = value['part']

                    # if key in accept_text_dict:
                    #     text_dict[key]["page"]=accept_text_dict[key]['page']
                    #     text_dict[key]["part"]=accept_text_dict[key]['part']



        finally:
            self.release_word(word, doc, doc_save=False, force_close=True)
        return revision_index

    def generate_accept_docx(self, doc_path):

        accept_doc_path = os.path.join(os.path.dirname(doc_path), "accept_" + os.path.basename(doc_path))
        logger.debug("accept_doc_path:%s" % accept_doc_path)
        shutil.copy(doc_path, accept_doc_path)
        # self.revision_stat(path_file_name=doc_path, save_file=accept_doc_path, accept=True)
        tmp_dir = tempfile.TemporaryDirectory()
        text_dict, _ = self.custom_doc_to_txt(path=accept_doc_path, txt_path=os.path.join(tmp_dir.name, "tmp.txt"),
                                              force_close=True)
        delete_file_if_exists(accept_doc_path)
        # delete_file_if_exists(tmp_dir.name)
        return text_dict

    def rebuild_page_and_add_comments_with_dict(self, doc_path, comments_dict, text_dict):
        if len(text_dict) == 0:
            return

        # accept_text_dict=self.generate_accept_docx(doc_path=doc_path)

        # 1. 记录diff.docx 有修订的index,
        # print(doc_path)
        revision_index = self.revisions_index_extract_and_build_page(doc_path=doc_path, comments_dict=comments_dict,
                                                                     text_dict=text_dict, accept_text_dict=None)
        # 2. Accept diff.docx,将这些index 的行内容提取从 accept_diff.docx出来
        # 3.  在与text_dict key 比对

        word, doc, e = self.doc_parse(doc_path)
        print("comments_dict", len(comments_dict), comments_dict.keys())
        print("revision_index", len(revision_index), revision_index.keys())
        print("text_dict", len(text_dict), text_dict.keys())
        try:
            for index, paragraph in enumerate(doc.Paragraphs):
                # 获取段落的修订信息
                if index in revision_index.keys():
                    # print("inhint",revision_index[index])
                    current_paragraph = doc.Paragraphs.Item(index + 1)
                    doc.Comments.Add(Range=current_paragraph.Range, Text=revision_index[index])
        finally:

            self.release_word(word, doc, force_close=True)
        return text_dict


if __name__ == "__main__":
    word_Mix_parse = WordMix()
    docpath="D:\\codes\\neolink-dataset\\contract-sentinel-web\\download_files\\蓝芯算力&世纪互联主合同-JH cmts-Vlegal1109.docx"
    txtpath="D:\\codes\\neolink-dataset\\contract-sentinel-web\\download_files\\蓝芯算力&世纪互联主合同-JH cmts-Vlegal1109.txt"

    try:
        word_Mix_parse.custom_doc_to_txt(path=docpath, txt_path=txtpath, force_close=True)
    except Exception as e:
        raise e
