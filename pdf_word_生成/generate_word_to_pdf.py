import subprocess
def convert_docx_to_pdf(docx_file, pdf_file):
    subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir', '/data/work/pydev/pdf生成', docx_file])

from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Inches
from docx.shared import Pt,Cm
from docx.oxml.ns import qn
from docx.shared import RGBColor
def create_custom_style(document, style_name, font_name, font_size, bold=False, italic=False, underline=False, color=None):
    # 创建自定义样式
    style = document.styles.add_style(style_name, 1)  # 1表示段落样式

    # 添加样式的字体
    style_font = style.font
    style_font.name = font_name
    style_font.size = Pt(font_size)
    style_font
    # style._element.rPr.rFonts.set(qn('w:eastAsia'), style.font)
    style_font.bold = bold
    style_font.italic = italic
    style_font.underline = underline
    if color:
        style_font.color.rgb = color

    return style

def apply_paragraph_style(paragraph, style_name):
    # 应用样式到段落
    paragraph.style = style_name

def create_word_document(output_path):
    # 创建一个新的 Word 文档
    doc = Document()

    # 设置页面的宽度和高度
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)  # A4 纸的高度（11.69英寸）

    # 设置页边距（示例为1英寸）
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # 创建自定义样式
    title_style = create_custom_style(doc, 'titleStyle', u'宋体', 26, bold=True, italic=False, underline=False,
                                       color=RGBColor(0, 0, 0))

    # 添加title并应用样式
    title_paragraph = doc.add_paragraph()
    title_run=title_paragraph.add_run('This is title')
    apply_paragraph_style(title_paragraph, title_style)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 换行
    title_run.add_break()


    normal_style = create_custom_style(doc, 'NormalStyle', u'宋体', 14, bold=False, italic=False, underline=False,
                                       color=RGBColor(0, 0, 0))

    normal_paragraph = doc.add_paragraph()
    normal_run=normal_paragraph.add_run('\t缩进是指段落与容器边缘(通常是页边距)之间的水平空间，段落可以在左右两边分别缩进。第一行缩进比其他部分缩进要大。第一行缩进较少用悬挂缩进。')
    apply_paragraph_style(normal_paragraph,normal_style)
    normal_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    normal_run.add_break()

    # 1. 标题
    # 插入标题
    doc.add_paragraph('标题', 'Title')
    # 插入一个副标题
    doc.add_paragraph('副标题', 'Subtitle')
    # 插入一级标题
    doc.add_paragraph('一级标题', style='Heading 1')
    # 插入二级标题
    doc.add_paragraph('二级标题', style='Heading 2')
    # 增加标题

    # 添加标题并应用字体格式
    title_paragraph = doc.add_paragraph()
    title_run=title_paragraph.add_run('Document Title')
    # 换行
    title_run.add_break()


    doc.add_paragraph(
        'This is a simple example of writing text and inserting a resized image into an A4-sized Word document.')
    doc.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
    # 插入图片并调整大小
    image_path = 'bar.png'  # 替换为你的图片路径
    a4_width = section.page_width - section.left_margin - section.right_margin
    a4_height = section.page_height - section.top_margin - section.bottom_margin

    # 插入调整后的图片到 Word 文档
    picture = doc.add_picture(image_path, width=a4_width)
    # png 是在段落里的，所以我们只需要对所在段落对齐即可
    doc.paragraphs[2].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 左对齐
    # 计算图片的缩放比例
    scale_factor =  picture.width / a4_width

    # 调整图片的大小
    picture.width = int(picture.width * scale_factor)
    picture.height = int(picture.height * scale_factor)

    # 保存 Word 文档
    doc.save(output_path)

    print(f"Word document created at: {output_path}")


# 输出 Word 文档路径
output_word_path = "output.docx"

# 创建 Word 文档并插入文字和调整大小后的图片
create_word_document(output_word_path)
convert_docx_to_pdf(output_word_path, 'output1.pdf')

