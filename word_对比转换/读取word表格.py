from docx import Document
import re, os


def read_docx_tables(docx_file):
    """
    读取 Word 文档中双列键值对型的表格
    :param docx_file: 文档文件名
    :return: [{key1: value1, ...}, ...]
    """
    # 创建文档对象，获得word文档
    document = Document(docx_file)
    # 读取表格集中的值
    doc_tables_values = []
    print(f"全部元素:{len(document.element.body.inner_content_elements)}")
    print(f"skip_table:{len(document.paragraphs)}")

    table_num=0
    for i, v in enumerate(document.element.body.inner_content_elements):
        if hasattr(v, "tblGrid") and hasattr(v, "tblPr"):
            table=document.tables[table_num]
            value_dict = {}
            for r in table.rows:
                cells = r.cells
                for ri in range(0,len(cells)):
                    k, v = cells[ri].text, cells[1].text
                    value_dict[k] = v
            doc_tables_values.append(value_dict)
    return doc_tables_values


path = os.path.dirname(__file__)
print(read_docx_tables(os.path.join(path, "my1.docx")))
