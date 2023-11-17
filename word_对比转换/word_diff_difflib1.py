import docx,os
import codecs
from difflib import HtmlDiff
path = "/data/work/pydev/word_对比转换/source_docx"
file1 = docx.Document(os.path.join(path,"bak","lock_sjhl_static_带宽罚则 锁定版(标准合同)_IDC主协议_北京世纪互联宽带数据中心托管服务协议(2023年版)-（预留机柜）.docx"))
file2 = docx.Document(os.path.join(path,"unlock_sjhl_static_带宽罚则 非锁定版(非标准合同)_IDC主协议_北京世纪互联宽带数据中心托管服务协议(2023年版)-（预留机柜）.docx"))

para1 = ''
para2 = ''

for para in file1.paragraphs:
    para1 = para1 + para.text + '\n'
    print(para1)
for para in file2.paragraphs:
    para2 = para2 + para.text + '\n'
    print(para2)

delta_html = HtmlDiff().make_file(para1.splitlines(), para2.splitlines())
with codecs.open('diff.html', 'w', encoding='utf-8') as f:
    f.write(delta_html)
