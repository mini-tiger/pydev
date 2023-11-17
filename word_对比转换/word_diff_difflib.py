# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""
import docx
import difflib
import os

'''在文件目录中存在两个待对比的word文档'''

# 获取文档对象
path = os.path.dirname(__file__)
path_file =  [f for f in os.listdir(path) if f.endswith('.docx')]

print('比较文档**--**%s**--**和文档**--**%s**--**的区别' % (path_file[0], path_file[1]))
file = docx.Document(path + "/" + path_file[0])
file2 = docx.Document(path + "/" + path_file[1])
print("%s共有---%s---个段落:" % (path_file[0], str(len(file.paragraphs))))
print("%s共有---%s---个段落:" % (path_file[1], str(len(file2.paragraphs))))
op = []
op2 = []

# 输出每一段的内容
for para in file.paragraphs:
    op.append(para.text)
print(op)
for para1 in file2.paragraphs:
    op2.append(para1.text)

diff = difflib.Differ()
numbe = 0
for d in range(len(op)):
    if op[d] != op2[d]:
        numbe += 1
        print('****第%s不同****' % (numbe))
        print('\n', path_file[0] + '的内容为：')
        print('    ~文档1：' + op[d])
        print(path_file[1] + '的内容为：')
        print('    ~文档2：' + op2[d], '\n')
        print('--------------------------------------------------------------------------------')
print('共有%s处不同' % (numbe))
print('对比完毕！！！！！！！！！！！！')