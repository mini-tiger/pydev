import docx, os

path = os.path.dirname(__file__)
# 读取第一个文档
doc1 = docx.Document(os.path.join(path, "my22.docx"))
doc1_text = '\n'.join([para.text for para in doc1.paragraphs])

# 读取第二个文档
doc2 = docx.Document(os.path.join(path, "my2.docx"))
doc2_text = '\n'.join([para.text for para in doc2.paragraphs])

import difflib

# 按行比较两个文档的内容，得到差异结果
differ = difflib.Differ()
diff_result = list(differ.compare(doc1_text.splitlines(), doc2_text.splitlines()))

# 输出差异结果
for line in diff_result:
    print(line)
