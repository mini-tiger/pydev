from docx import Document
from docx.shared import Pt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
# 创建一个新的文档
doc = Document()

# 定义一个方法来创建没有符号的标题
def set_font(run, name, size: int, isBole: bool = False, isItalic: bool = False, isUnderline: bool = False,
             color: RGBColor = None):
    # doc.styles['Normal'].font
    run.font.name = f'{name}'  # 款式
    run._element.rPr.rFonts.set(qn('w:eastAsia'), f'{name}')
    run.font.size = Pt(int(f'{size}'))  # 大小
    run.font.bold = isBole  # 加粗
    run.font.italic = isItalic  # 倾斜
    run.font.underline = isUnderline  # 下划线
    if color is not None:
        run.font.color.rgb = RGBColor(255, 0, 0)  # 颜色
def center_line(doc, text):
    pg = doc.add_paragraph()
    pg.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平对齐设为居中
    return pg.add_run(text)
def add_custom_heading(doc, text, level):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(14 + (4 - level) * 2)  # 根据标题级别设置字体大小
    paragraph.style = doc.styles['Heading %d' % level]

# 添加标题
run1 = center_line(doc, '主标题')
set_font(run1, '黑体', 38, True)
# 标记目录的位置
catalog_p = doc.add_paragraph('')
add_custom_heading(doc, '副标题', 2)
add_custom_heading(doc, '三级标题', 3)

# 添加一些正文内容
doc.add_paragraph('这是正文内容的一部分。')
# 添加节
doc.add_section()
# catalog_doc(doc, catalog_p)

# 保存文档
doc.save('/mnt/191/example.docx')
