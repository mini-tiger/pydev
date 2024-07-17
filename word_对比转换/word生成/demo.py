# -*- coding: utf-8 -*-

import json
import os.path
import re
import time

import uvicorn
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from fastapi import FastAPI, Request
from starlette.middleware.cors import CORSMiddleware
from starlette.responses import FileResponse

from utils.Logger import logger

app = FastAPI()
origins = ['*']
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# logger = configure_logging()

@app.post("/generator")
async def generator(items: dict, request: Request = None):
    url = ''
    if request is not None:
        url = request.url
    logger.info(f"start---{url}---{items.get('title')}--{items.get('config_file')}")
    path = items.get("path")
    title = items.get('title')
    year = items.get('year')
    author = items.get('author')
    config_file = 'tpl-中咨'
    if 'config_file' in items:
        config_file = items.get("config_file")
    property = get_config(f'{config_file}.properties')
    property['icon_name'] = items.get('icon_name')
    property['company'] = items.get('company')
    property['copyright'] = items.get('copyright')

    if isinstance(items.get('Sections'), str):
        sections = json.loads(str(items.get('Sections')).replace('\'', '\"'))
    else:
        sections = items.get('Sections')
    sections.sort(key=lambda x: x['index'])

    doc = Document()
    # 初始化
    init_content(doc, title, property)
    # 开始写入文档
    line = center_line(doc, '目录')
    set_font(line, '黑体', 18, True)

    # 标记目录的位置
    # catalog_p = doc.add_paragraph('')
    # 添加节
    doc.add_section()
    for item in sections:
        iterate_over_arrays(doc, 1, item, [0])

    style = doc.styles['Normal']
    style.font.name = '宋体'  # 必须先设置font.name
    style.font.size = Pt(11)  # 必须先设置font.name
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.color.rgb = RGBColor(0, 0, 0)  # 颜色

    # 页脚页码
    insert_page_number(doc)
    # 文件名
    file_name = generate_file_name(title)
    # print(f'--------------{name}-----------------')
    doc.save(file_name)
    update_toc(os.path.abspath(file_name))
    file_path, full_name = os.path.split(file_name)
    f_name, ext = os.path.splitext(full_name)
    print(file_name)
    logger.info(f"end---{items.get('title')}---{file_name}")
    return FileResponse(file_name, filename=full_name)


def init_content(doc, title, property):
    icon = property.get('icon_name')
    company = property.get('company')
    copyright = property.get('copyright')

    # 封面-cover
    icon = os.path.abspath(f'./config/icon/{icon}')
    # pip install Pillow
    from PIL import Image
    img = Image.open(icon)
    # imgSize = img.size  # 大小/尺寸
    # w = img.width  # 图片的宽
    # h = img.height  # 图片的高
    # f = img.format  # 图像格式
    # print(imgSize)
    # print(w, h, f)

    doc.add_picture(icon, width=Inches(int(img.width / 80 + 1)))
    blank_line(doc, 3)

    run1 = center_line(doc, title)
    set_font(run1, '黑体', 38, True)
    run2 = center_line(doc, f"（{time.strftime('%Y年')}）")
    set_font(run2, '黑体', 18, True)

    blank_line(doc, 8)
    run3 = center_line(doc, company)
    set_font(run3, '仿宋', 18, True)
    run3 = center_line(doc, time.strftime('%Y年%m月%d日'))
    set_font(run3, '仿宋', 18, True)

    cover = doc.add_paragraph('')
    break_page(cover)
    if not copyright is None:
        add_copyright(doc, copyright)


def add_copyright(doc, copyright):
    # 版权声明 - Copyright
    blank_line(doc, 3)
    add_under_line(doc)
    blank_line(doc, 1)
    run4 = center_line(doc, '版权声明')
    set_font(run4, '仿宋', 18, True)
    # <w:bCs/>
    # <w:spacing w:val="59"/>
    # <w:w w:val="141"/>
    # <w:szCs w:val="36"/>
    run4._element.rPr.append(OxmlElement('w:bCs'))
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), '59')
    w = OxmlElement('w:w')
    w.set(qn('w:val'), '141')
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '36')
    run4._element.rPr.append(spacing)
    run4._element.rPr.append(w)
    run4._element.rPr.append(szCs)
    # blank_line(doc, 1)
    # =======================================================
    add_under_line(doc)
    # ==================================================
    blank_line(doc, 1)
    pg = doc.add_paragraph()
    pg.paragraph_format.line_spacing = 1  # 行间距，1.5倍行距
    pg.paragraph_format.first_line_indent = Pt(36)  # 首行缩进10磅
    run4 = pg.add_run(copyright)
    set_font(run4, '宋体', 14)

    cover = doc.add_paragraph('')
    break_page(cover)
    # 前言-preface
    pass


def blank_line(doc, n: int):
    for i in range(n):
        doc.add_paragraph().paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平对齐设为居中


def center_line(doc, text):
    pg = doc.add_paragraph()
    pg.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平对齐设为居中
    return pg.add_run(text)


def set_font(run, name, size: int, isBole: bool = False, isItalic: bool = False, isUnderline: bool = False,
             color: RGBColor = None):
    # doc.styles['Normal'].font
    run.font.name = f'{name}'  # 款式
    run._element.rPr.rFonts.set(qn('w:eastAsia'), f'{name}')
    run.font.size = Pt(int(f'{size}'))  # 大小
    run.font.bold = isBole  # 加粗
    run.font.italic = isItalic  # 倾斜
    run.font.underline = isUnderline  # 下划线
    if color is not None:
        run.font.color.rgb = RGBColor(255, 0, 0)  # 颜色


def iterate_over_arrays(doc, idx: int, item, count):
    count[0] = count[0] + 1
    # text = str(item.get('index')) + ' ' + item.get('Section')
    text = item.get('Section')
    add_title_with_bookmark(doc, text, style=f'Heading {idx}', bookmark_id=count[0])
    if 'data' in item:
        add_doc_data(doc, item.get('data'))

        source = f"{item.get('context_source')}"
        if source.startswith('\n'):
            source = source[1:]
        # print(source + '---' + str(len(source)))
        if len(source) > 10:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.5  # 行间距，1.5倍行距
            paragraph.paragraph_format.first_line_indent = Pt(22)  # 首行缩进10磅
            paragraph.paragraph_format.space_before = 0  # 段前30磅
            paragraph.paragraph_format.space_after = Pt(15)  # 段后15磅
            run = paragraph.add_run(f'来源：“{source}”。')
            run.font.italic = True

    if 'child' in item:
        item.get('child').sort(key=lambda x: x['index'])
        for i, cc in enumerate(item.get('child')):
            iterate_over_arrays(doc, idx + 1, cc, count)


def generate_file_name(title):
    UTC_FORMAT = "%Y%m%dT%H%M%S"
    utcTime = time.strftime(UTC_FORMAT, time.localtime())
    import random
    return f'./attachments/{title}-{utcTime}-{str(random.randint(1000, 9999))}.docx'


# def init_items(items: dict):
#     items["path"] = './'
#     items['title'] = '大数据白皮书'
#     items['year'] = 2024
#     items['author'] = '中国信息通信研究院'
#
#     if len(items.get('title')) < 8:
#         items['title'] = re.sub(r"(?<=\w)(?=(?:\w\w)+$)", " ", items.get('title'))


# 添加标题和书签
def add_title_with_bookmark(doc, text, style, bookmark_id):
    paragraph = doc.add_paragraph(text, style=style)
    paragraph.style.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.style.font.name = u'宋体'  # 款式
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), u'宋体')
    rPr = OxmlElement('w:rPr')
    rPr.append(rFonts)
    paragraph._element.r_lst[0].insert(0, rPr)

    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:afterLines'), '50')
    spacing.set(qn('w:after'), '120')
    paragraph._element.pPr.insert(1, spacing)

    # paragraph.style.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run = paragraph.add_run()

    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), text)
    tag.append(start)

    tr = OxmlElement('w:r')
    tr.text = ''
    tag.append(tr)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    end.set(qn('w:name'), text)
    tag.append(end)


def break_page(paragraph):
    # 特定段落分页
    break_page_p = paragraph.insert_paragraph_before()
    break_page_p.add_run().add_break(WD_BREAK.PAGE)


def catalog_doc(doc, catalog_p):
    # 自动生成目录内容，不包含页码
    for paragraph in doc.paragraphs:
        if 'Heading' in paragraph.style.name:
            b = paragraph._element.findall('.//' + qn('w:bookmarkStart'))
            bookmark_name = b[0].get(qn('w:name'))
            text = paragraph.text
            level = int(paragraph.style.name[-1])
            # new_p = doc.add_paragraph('text')
            # print(text, bookmark_name)

            toc_paragraph = catalog_p.insert_paragraph_before(style='Normal')
            # 二级标题设置缩进
            if level == 2:
                toc_paragraph.paragraph_format.first_line_indent = Pt(24)

            # 设置制表符，可显示页码前的"…………"
            tabs = OxmlElement('w:tabs')
            tab1 = OxmlElement('w:tab')
            tab1.set(qn('w:val'), "left")
            tab1.set(qn('w:leader'), "hyphen")  # hyphen,dot
            tab1.set(qn('w:pos'), "8400")
            tabs.append(tab1)
            toc_paragraph._p.pPr.append(tabs)
            # toc_paragraph若未设定style，toc_paragraph._p没有pPr属性，需注释前一句代码，使用以下语句
            # pPr = OxmlElement('w:pPr')
            # pPr.append(tabs)
            # toc_paragraph._p.append(pPr)

            # hyperlink = OxmlElement('w:hyperlink')
            # hyperlink.set(qn('w:anchor'), bookmark_name)
            # hyperlink.set(qn('w:history'), '1')

            hr1 = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), "a4")
            rPr.append(rStyle)
            hr1.text = text
            hr1.append(rPr)
            # hyperlink.append(hr1)
            toc_paragraph._p.append(hr1)

            # hr2 = OxmlElement('w:r')
            # rPr = OxmlElement('w:rPr')
            # fldChar = OxmlElement('w:fldChar')
            # fldChar.set(qn('w:fldCharType'), 'begin')
            # hr2.append(rPr)
            # hr2.append(fldChar)
            # # hyperlink.append(hr2)
            # toc_paragraph._p.append(hr2)

            # hr3 = OxmlElement('w:r')
            # rPr = OxmlElement('w:rPr')
            # instrText = OxmlElement('w:instrText')
            # instrText.set(qn('xml:space'), 'preserve')
            # instrText.text = ' PAGEREF {} \h '.format(bookmark_name)
            # hr3.append(rPr)
            # hr3.append(instrText)
            # # hyperlink.append(hr3)
            # toc_paragraph._p.append(hr3)

            # hr4 = OxmlElement('w:r')
            # rPr = OxmlElement('w:rPr')
            # fldChar = OxmlElement('w:fldChar')
            # fldChar.set(qn('w:fldCharType'), 'separate')
            # hr4.append(rPr)
            # hr4.append(fldChar)
            # # hyperlink.append(hr4)
            # toc_paragraph._p.append(hr4)

            hrt = OxmlElement('w:r')
            tab = OxmlElement('w:tab')
            hrt.append(tab)
            # hyperlink.append(hrt)
            toc_paragraph._p.append(hrt)

            hr5 = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            hr5.text = ''
            hr5.append(rPr)
            # hyperlink.append(hr5)
            toc_paragraph._p.append(hr5)

            # hr6 = OxmlElement('w:r')
            # rPr = OxmlElement('w:rPr')
            # fldChar = OxmlElement('w:fldChar')
            # fldChar.set(qn('w:fldCharType'), 'end')
            # hr6.append(rPr)
            # hr6.append(fldChar)
            # # hyperlink.append(hr6)
            # toc_paragraph._p.append(hr6)
            # toc_paragraph._p.append(hyperlink)
            # hr7 = OxmlElement('w:r')
            # t7 = OxmlElement('w:t')
            # t7.text = '1'
            # hr7.append(t7)
            # toc_paragraph._p.append(hr7)


def add_under_line(doc):
    paragraph = doc.add_paragraph()
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')  # top
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    paragraph.paragraph_format.line_spacing = 1  # 行间距，1.5倍行距
    paragraph.paragraph_format.line_spacing = 0  # 行间距，固定值20磅
    paragraph.paragraph_format.space_before = 0  # 段前30磅
    paragraph.paragraph_format.space_after = 0  # 段后15磅
    pass


def add_footer_number(run):
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'Page'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)


def insert_page_number(doc):
    footer = doc.sections[1].footer  # 获取第一个节的页脚
    footer.is_linked_to_previous = False  # 编号续前一节
    paragraph = footer.paragraphs[0]  # 获取页脚的第一个段落
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 页脚居中对齐
    run_footer = paragraph.add_run()  # 添加页脚内容
    add_footer_number(run_footer)
    font = run_footer.font
    font.name = 'Times New Roman'  # 新罗马字体
    font.size = Pt(10)  # 10号字体
    font.bold = True  # 加粗

    sectPr = doc.sections[1]._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    # pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)


# def covert_md_to_word(md_str):
#     pattern = re.compile(r'[\n]+[1-9]\d*.[ ][\*]*[\u4e00-\u9fa5]')
#     r_list = pattern.findall(md_str)
#     for r in r_list:
#         md_str = md_str.replace(r, r.replace(' ', ''))
#
#     import pypandoc
#     import uuid
#     file_dirs = os.path.abspath("./attachments/tmp")
#     file_name = f'{file_dirs}//h2w-%s.docx' % str(uuid.uuid1()).replace('-', '')
#     html_str = pypandoc.convert_text(md_str, 'html', format='md')
#     pypandoc.convert_text(html_str, 'docx', format='html', outputfile=file_name)
#     doc = Document(file_name)
#     os.remove(file_name)
#     for para in doc.paragraphs:
#         para.paragraph_format.line_spacing = 1.5  # 行间距，1.5倍行距
#         para.paragraph_format.first_line_indent = Pt(22)  # 首行缩进10磅
#     return doc


def get_config(config_file):
    tpl_dirs = os.path.abspath(f'./config/{config_file}')
    from utils.property import Properties
    return Properties(tpl_dirs).getProperties()


def update_toc(docx_file):  # word路径
    import win32com.client
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = 0  # 设置应用可见
    word.DisplayAlerts = 0
    doc = word.Documents.Open(docx_file)  # 使用微软office打开word
    toc_count = doc.TablesOfContents.Count  # 判断是否有无目录，如果数量是1则代表已经有目录了
    if toc_count == 0:
        for i, p in enumerate(doc.Paragraphs):  # 遍历word中的内容
            if '目录' in p.Range.Text:  # 用于指定目录页面，看下面提示
                p.Range.InsertParagraphAfter()  # 添加新的段落
                p.Range.InsertAfter("---")
                parag_range = doc.Paragraphs(i + 2).Range
                doc.TablesOfContents.Add(Range=parag_range, UseHeadingStyles=True, LowerHeadingLevel=3,
                                         UseHyperlinks=True)  # 生成目录对象
    elif toc_count == 1:
        toc = doc.TablesOfContents(1)
        toc.Update()
    doc.Close(SaveChanges=True)
    word.Quit()


def delete_files(directory):
    file_list = os.listdir(directory)
    for file in file_list:
        file_path = os.path.join(directory, file)
        if os.path.isfile(file_path):
            os.remove(file_path)


def add_doc_data(doc, data):
    # ------------------------------------------------
    text = data.replace('**', '').replace('\n\n', '\n')
    if text.startswith('\n'):
        text = text[1:]
    content = text.split('\n')
    for c in content:
        paragraph = doc.add_paragraph(c)
        # paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 行距固定值
        # paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # 多倍行距
        paragraph.paragraph_format.line_spacing = 1.5  # 行间距，1.5倍行距
        # paragraph.paragraph_format.line_spacing = Pt(0)  # 行间距，固定值20磅
        paragraph.paragraph_format.first_line_indent = Pt(22)  # 首行缩进10磅
        paragraph.paragraph_format.space_before = 0  # 段前30磅
        paragraph.paragraph_format.space_after = 0  # 段后15磅
    # # ====================
    # if not data.startswith('\n'):
    #     data = '\n' + data
    # doc_md = covert_md_to_word(data)
    # for element in doc_md.element.body:  # 拷贝文件中的信息，# 追加第二个文档内容到第一个文档末尾
    #     if isinstance(element, CT_P):
    #         # doc.element.body.append(element)
    #         doc.element.body.insert(-1, element)


# ======================================================
def log_loop():
    import time
    interval = 10
    while True:
        # 将日期时间对象格式化为字符串
        utcTime = time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime())
        logger.info(f'打印日志-{utcTime}')
        time.sleep(interval)


if __name__ == '__main__':
    # # 创建多个线程来调用日志打印函数
    # import threading
    # thread = threading.Thread(target=log_loop, name=f"Thread-1", daemon=True)
    # thread.start()

    # log_level = 'info'
    uvicorn.run('generate_doc:app', host='0.0.0.0', port=7891, reload=True)
