from docx import Document
from docx.shared import Pt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import nsdecls
import docx

# 创建一个新的文档

doc = Document()


# 定义一个方法来创建没有符号的标题
def set_font(run, name, size: int, isBole: bool = False, isItalic: bool = False, isUnderline: bool = False,
             color: RGBColor = None):
    # doc.styles['Normal'].font
    run.font.name = f'{name}'  # 款式
    run._element.rPr.rFonts.set(qn('w:eastAsia'), f'{name}')
    run.font.size = Pt(int(f'{size}'))  # 大小
    run.font.bold = isBole  # 加粗
    run.font.italic = isItalic  # 倾斜
    run.font.underline = isUnderline  # 下划线
    if color is not None:
        run.font.color.rgb = RGBColor(255, 0, 0)  # 颜色


def center_line(doc, text):
    pg = doc.add_paragraph()
    pg.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平对齐设为居中
    return pg.add_run(text)


def add_custom_heading(doc, text, level):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(14 + (4 - level) * 2)  # 根据标题级别设置字体大小
    paragraph.style = doc.styles['Heading %d' % level]


def insert_page_number1(doc):
    # 找到页码需要插入的位置，通常是页脚
    footer = doc.sections[0].footer
    # 找到页码需要插入的段落
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 页脚居中对齐
    # 创建页码对象
    page_number = docx.oxml.shared.OxmlElement("w:fldSimple")
    page_number.set(qn("w:instr"), r'PAGE')

    # 创建域对象
    run = docx.oxml.shared.OxmlElement("w:r")
    run.append(page_number)

    # 将域对象插入到段落中
    paragraph._p.append(run)


def insert_page_number(doc):
    footer = doc.sections[1].footer  # 获取第一个节的页脚
    footer.is_linked_to_previous = False  # 编号续前一节
    paragraph = footer.paragraphs[0]  # 获取页脚的第一个段落
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 页脚居中对齐
    run_footer = paragraph.add_run()  # 添加页脚内容
    add_footer_number(run_footer)
    font = run_footer.font
    font.name = 'Times New Roman'  # 新罗马字体
    font.size = Pt(10)  # 10号字体
    font.bold = True  # 加粗

    sectPr = doc.sections[1]._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    # pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)


def add_footer_number(run):
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'Page'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)


# 添加标题
run1 = center_line(doc, '主标题')
set_font(run1, '黑体', 38, True)

# 标记第一页的位置
doc.add_section()

# 标记目录的位置
catalog_p = doc.add_paragraph('')
add_custom_heading(doc, '副标题', 2)
add_custom_heading(doc, '三级标题', 3)

# 添加一些正文内容
doc.add_paragraph('这是正文内容的一部分。')
doc.add_page_break()

# 页码
# https://blog.csdn.net/Adam_captain/article/details/134728397
# insert_page_number(doc)
insert_page_number1(doc)

paragraphs = doc.paragraphs
print(paragraphs[3]._p.xml)
print(paragraphs[3].style.element.xml)

link_id = paragraphs[3].style.element.xpath('string(.//w:name/@w:val)')
# font_name_zh = doc.styles.element.xpath(f'string(.//w:style[@w:styleId={link_id}]//w:rFonts/@w:eastAsia)')
# font_name = doc.styles.element.xpath(f'string(.//w:style[@w:styleId={link_id}]//w:rFonts/@w:ascii)')
print(link_id)
# print(font_name_zh)
# print(font_name)
# 保存文档
doc.save('/mnt/191/example.docx')

